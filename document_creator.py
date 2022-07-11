import pandas as pd
import numpy as np
import pdb
from docx import Document
from docx.shared import Mm
from collections import defaultdict

from report_template import *



'''Create word document using word'''
Anomali_report_full = Document()

section = Anomali_report_full.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)


'''Mechanical Availability'''


'''
formatting MA table
'''

table_MA = Anomali_report_full.add_table(rows=1,cols=3)
table_MA.style = 'Table Grid'
hdr_cells_MA = table_MA.rows[0].cells
hdr_cells_MA[0].text = 'Fleet/Model'
hdr_cells_MA[1].text = 'MA this week'
hdr_cells_MA[2].text = 'Comments on Non-Compliance'

for i in MA_report:
    
    TOU_MA = i.get('Total Unit')
    PMD_MA = i.get('PMD total hours')
    UMD_MA = i.get('UMD total hours')
    PDAM_MA = i.get('PDAM total hours')
    
    row_cells_MA = table_MA.add_row().cells
    row_cells_MA[0].text = i.get('Fleet')
    row_cells_MA[1].text = "{:.2%}".format(i.get('MA'))
    
    sch = i.get('Scheduled maintenance')
    Unsch = i.get('Unscheduled maintenance')
    
    Scheduled_maintenance = '\n\n'.join([c for c in sch[1:]])
    Unscheduled_maintenance = '\n\n'.join([c for c in Unsch[1:]])
    
    non_compliance = f'Number of Units = {TOU_MA} \nPMD Total = {PMD_MA} \nUMD Total= {UMD_MA} \nPDAM Total hours= {PDAM_MA}\n Scheduled maintenance :\n{Scheduled_maintenance}\nUnscheduled maintenance :\n{Unscheduled_maintenance}'

    row_cells_MA[2].text = non_compliance

Anomali_report_full.add_page_break()
    
'''
Mean Time Between Failure
'''

'''
formatting MTBF table 
'''

table_MTBF = Anomali_report_full.add_table(rows=1, cols=3)
table_MTBF.style = 'Table Grid'

hdr_cells_MTBF = table_MTBF.rows[0].cells
hdr_cells_MTBF[0].text = 'Fleet/Model '
hdr_cells_MTBF[1].text = 'MTBF this week'
hdr_cells_MTBF[2].text = 'Comments on Non-Compliance'

for d in MTBF_report:
    
    
    
    TOU = d.get('Total Unit')
    THEON = d.get('Total Hours Engine ON')
    COF = d.get('Count of failures')
    
    
    
    
    row_cells = table_MTBF.add_row().cells
    row_cells[0].text = d.get('Fleet')
    row_cells[1].text = str(d.get('MTBF'))
    
    gp_components = []
    
    
    for system in d.get('Failures by Comp Gp with repetitive issue'):
            for key, value in system.items():
                s = ", ".join(value)
                ter = f'{key}: {s}'
                gp_components.append(ter)
            
            
    f_list = '\n'.join([d for d in gp_components[0:]])
    gp_comp_list = f'Number of Units = {TOU} \nTotal Hours Engine On = {THEON} \n Count of failures = {COF} \n Failures by Comp Gp with repetitive issue : \n {f_list}'
    
    row_cells[2].text = gp_comp_list

Anomali_report_full.add_page_break()
    
'''
Mean Time to Repair Failure
'''

'''
formatting MTTRF table 
'''
table_MTTR = Anomali_report_full.add_table(rows=1, cols=3)
table_MTTR.style = 'Table Grid'

hdr_cells_MTTR = table_MTTR.rows[0].cells
hdr_cells_MTTR[0].text = 'Fleet/Model '
hdr_cells_MTTR[1].text = 'MTTR this week'
hdr_cells_MTTR[2].text = 'Comments on Non-Compliance'

for i in MTTR_report:
    
    TOU_MTTRF = i.get('MTTR')
    Hours_Fail_MTTRF = i.get('Hours Maintenance with Failure is Yes')
    COF_MTTRF = i.get('Count of Failures')
    DBF_MTTRF = i.get('Detail Breakdown with Failure is Yes')
    detail_breakdown_MTTRF = '\n'.join([c for c in DBF_MTTRF[1:]])
        
    comment_MTTR = f'Number of Units = {TOU} \nHours maintenance with is Failure Yes ={Hours_Fail_MTTRF} \nCount of failures/Is Failure Yes = {COF_MTTRF} \n\n\n Detail breakdown with failure is yes:\n{detail_breakdown_MTTRF}'
    
    row_cells = table_MTTR.add_row().cells
    row_cells[0].text = i.get('Fleet')
    row_cells[1].text = str(i.get('MTTR'))
    row_cells[2].text = comment_MTTR

Anomali_report_full.save('Anomali_trial1.docx')
del Anomali_report_full